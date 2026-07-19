import io
import os
import textwrap
import pandas as pd 
import matplotlib
matplotlib.use('Agg')
import matplotlib.pyplot as plt
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from fpdf import FPDF
from classification import QuestionInfo, QuestionType
from summary import QuestionSummary
from typing import List
from lang import LANGUAGES

CHART_DPI = 150
FONT_SIZE_CHART = 10
BAR_WIDTH = 0.6

FONT_FILENAME = "Tinos-Regular.ttf"
FONT_PATH = os.path.join(os.getcwd(), FONT_FILENAME)
FONT_URL = "https://github.com/google/fonts/raw/main/apache/tinos/Tinos-Regular.ttf"


def set_table_borders(table):
    tbl = table._tbl
    tblPr = tbl.tblPr
    tblBorders = OxmlElement('w:tblBorders')
    for border_name in ['top', 'left', 'bottom', 'right', 'insideH', 'insideV']:
        border = OxmlElement(f'w:{border_name}')
        border.set(qn('w:val'), 'single')
        border.set(qn('w:sz'), '4')
        border.set(qn('w:space'), '0')
        border.set(qn('w:color'), '000000')
        tblBorders.append(border)
    tblPr.append(tblBorders)

class PDFReport(FPDF):
    def __init__(self, selected_lang="UA"):
        super().__init__()
        self.selected_lang = selected_lang

        self.add_font("ArialUni", "", "arial.ttf")         # Звичайний
        self.add_font("ArialUni", "B", "arialbd.ttf")


    def header(self):
        if self.page_no() > 1:
            try:
                self.set_font("ArialUni", "B", size=12)
                self.cell(0, 10, self.lang_dict['report_title'], ln=1, align='R')
            except:
                self.set_font("ArialUni", "B", 10)
                self.cell(0, 10, self.lang_dict['report_title'], ln=1, align='R')
        else:
            self.ln(5)
        pass

    def footer(self):
        self.set_y(-15)
        try:
            self.set_font("ArialUni", size=8)
        except:
            self.set_font("ArialUni", "I", 8)
        self.cell(0, 10, f'{self.page_no()}', align='C')

def create_chart_image(qs: QuestionSummary, selected_lang="UA") -> io.BytesIO:
    plt.close('all')
    plt.clf()
    plt.rcParams.update({'font.size': FONT_SIZE_CHART})
    
    lang_dict = LANGUAGES[selected_lang]

    col_variant = qs.table.columns[0]
    col_count = qs.table.columns[1]

    labels = qs.table[col_variant].astype(str).tolist()

    values = pd.to_numeric(qs.table[col_count], errors='coerce').fillna(0)
    if values.sum() == 0:
        return None
    
    wrapped_labels = [textwrap.fill(l, 25) for l in labels]

    is_scale = (qs.question.qtype == QuestionType.SCALE)
    if not is_scale:
        try:
            vals = pd.to_numeric(qs.table[col_variant], errors='coerce')
            if vals.notna().all() and vals.min() >= 0 and vals.max() <= 10:
                is_scale = True
        except: pass

    if is_scale:
        # СТОВПЧИКОВА
        fig = plt.figure(figsize=(6.0, 4.0))
        bars = plt.bar(wrapped_labels, values, color='#4F81BD', width=BAR_WIDTH)
        plt.ylabel(lang_dict["count"])
        plt.grid(axis='y', linestyle='--', alpha=0.5)
        for bar in bars:
            height = bar.get_height()
            plt.text(bar.get_x() + bar.get_width()/2., height + 0.1,
                     f'{int(height)}', ha='center', va='bottom', fontweight='bold')
    else:
        # КРУГОВА
        fig = plt.figure(figsize=(6.0, 4.0))
        colors = ['#4F81BD', '#C0504D', '#9BBB59', '#8064A2', '#4BACC6', '#F79646']
        c_arg = colors[:len(values)] if len(values) <= len(colors) else None
        
        wedges, texts, autotexts = plt.pie(
            values, labels=None, autopct='%1.1f%%', startangle=90,
            pctdistance=0.8, colors=c_arg, radius=1.0,
            textprops={'fontsize': FONT_SIZE_CHART}
        )
        for autotext in autotexts:
            autotext.set_color('white')
            autotext.set_weight('bold')
            import matplotlib.patheffects as path_effects
            autotext.set_path_effects([path_effects.withStroke(linewidth=2, foreground='#333333')])

        plt.axis('equal')
        cols = 2 if len(labels) > 3 else 1
        plt.legend(wrapped_labels, loc="upper center", bbox_to_anchor=(0.5, 0.0), ncol=cols, frameon=False, fontsize=9)

    plt.tight_layout()
    img_stream = io.BytesIO()
    plt.savefig(img_stream, format='png', dpi=CHART_DPI, bbox_inches='tight')
    plt.close(fig)
    img_stream.seek(0)
    return img_stream

def build_docx_report(original_df, sliced_df, summaries, range_info, selected_lang="UA") -> bytes:
    lang = LANGUAGES[selected_lang]
    doc = Document()
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Times New Roman'
    font.size = Pt(12)

    head = doc.add_heading(lang['report_title'], 0)
    head.alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph(f"{lang['total_surveys']}: {len(original_df)}")
    doc.add_paragraph(f"{lang['processed_surveys']}: {len(sliced_df)}")
    doc.add_paragraph(f"{lang['range']}: {range_info}")
    'doc.add_page_break()'

    for qs in summaries:
        if qs.table.empty: continue
        
        p = doc.add_paragraph()
        runner = p.add_run(f"{qs.question.code}. {qs.question.text}")
        runner.bold = True
        runner.font.size = Pt(14)
        
        table = doc.add_table(rows=1, cols=3)
        set_table_borders(table)
        hdr = table.rows[0].cells
        hdr[0].text = lang["variant"]; hdr[1].text = lang["count"]; hdr[2].text = '%'
        
        for row in qs.table.itertuples(index=False):
            rc = table.add_row().cells
            rc[0].text = str(row[0])
            rc[1].text = str(row[1])
            rc[2].text = str(row[2])

        try:
            img_stream = create_chart_image(qs, selected_lang)
            if img_stream is not None:
                doc.add_picture(img_stream, width=Inches(5.5))
                doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
        except Exception as e:
            print(f"Error creating chart for question {qs.question.code}: {e}")
        doc.add_paragraph("\n")
    
    doc.add_paragraph()
    doc.add_paragraph()

    footer_text = [
        lang["report_footer"],
        lang["report_footer2"],
    ]

    for line in footer_text:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT 
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(2)
        
        run = p.add_run(line)
        run.font.name = 'Times New Roman'
        run.font.size = Pt(12)                   
        run.font.color.rgb = RGBColor(100, 100, 100) 

    output = io.BytesIO()
    doc.save(output)
    return output.getvalue()